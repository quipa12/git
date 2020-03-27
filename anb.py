from docx import Document
import glob


lista1 = glob.glob("files/*.docx")
print(lista1)
ind1 = 0
for file in lista1:
    doc = Document(file)
    print(doc)
    lista_file_parse = file.split('.')
    file_edit = lista_file_parse[0]
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    '\n'.join(fullText)

    for para in fullText:
        if 'bananas' in para:
            para1 = 'Eu quero maçãs'
            ind = fullText.index(para)
            fullText.remove(para)
            fullText.insert(ind,para1)

    document2 = Document()
    for line in fullText:
        document2.add_paragraph(line)
    document2.save('destino\\'+file_edit+'atualizado'+'.docx')
    ind1 = ind1 + 1
str(input('Pressione qualquer tecla pra sair '))
